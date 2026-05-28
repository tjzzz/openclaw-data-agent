#!/usr/bin/env python3
"""
AI Humanizer - Web Application
Flask app that provides AI detection and text humanization services.
"""

import os
import re
import uuid
import json
import io
from functools import wraps
from datetime import datetime, timedelta
from flask import Flask, render_template, request, jsonify, session, send_file, g
from flask.sessions import SessionInterface, SessionMixin
from uuid import uuid4
import pickle as _pickle
from dotenv import load_dotenv

load_dotenv()

from ai_checker import analyze_text, analyze_by_paragraphs
from humanize import humanize_text
from models import init_db, get_connection, User, Order
from payment_adapter import create_payment_adapter
from humanizer_adapter import RuleBasedHumanizer

app = Flask(__name__)
if not os.environ.get('SECRET_KEY'):
    raise RuntimeError("SECRET_KEY environment variable must be set. Generate one with: python -c 'import secrets; print(secrets.token_hex(32))'")
app.secret_key = os.environ['SECRET_KEY']
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['PAYMENT_ADAPTER'] = os.environ.get('PAYMENT_ADAPTER', 'mock')
app.config['HUMANIZER_ADAPTER'] = 'rule_based'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=2)

# Server-side filesystem session (avoids Flask's 4KB cookie limit)
_session_dir = os.path.join(os.path.dirname(__file__), 'instance', 'flask_session')
os.makedirs(_session_dir, exist_ok=True)


class _FileSystemSession(dict, SessionMixin):
    """Minimal server-side session stored on the filesystem."""
    def __init__(self, data=None, sid=None, new=False):
        super().__init__(data or {})
        self.sid = sid or uuid4().hex
        self.new = new
        self.modified = True


class _FileSystemSessionInterface(SessionInterface):
    """Filesystem-based session interface using pickle serialization."""
    def __init__(self, session_dir):
        self.session_dir = session_dir
        os.makedirs(session_dir, exist_ok=True)

    def open_session(self, app, request):
        _cookie_name = app.config.get('SESSION_COOKIE_NAME', 'session')
        try:
            sid = request.cookies.get(_cookie_name)
        except Exception:
            sid = None
        if sid:
            path = os.path.join(self.session_dir, sid)
            if os.path.exists(path):
                try:
                    with open(path, 'rb') as f:
                        data = _pickle.load(f)
                    return _FileSystemSession(data=data, sid=sid)
                except Exception:
                    pass
        return _FileSystemSession(new=True)

    def save_session(self, app, session, response):
        _cookie_name = app.config.get('SESSION_COOKIE_NAME', 'session')
        if not session or not hasattr(session, 'sid'):
            if session and hasattr(session, 'sid'):
                path = os.path.join(self.session_dir, session.sid)
                if os.path.exists(path):
                    os.remove(path)
            response.delete_cookie(_cookie_name)
            return

        path = os.path.join(self.session_dir, session.sid)
        with open(path, 'wb') as f:
            _pickle.dump(dict(session), f)

        response.set_cookie(
            _cookie_name, session.sid,
            max_age=self.get_expiration_time(app, session),
            httponly=self.get_cookie_httponly(app),
            domain=self.get_cookie_domain(app),
            path=self.get_cookie_path(app),
            secure=self.get_cookie_secure(app),
            samesite=self.get_cookie_samesite(app),
        )


app.session_interface = _FileSystemSessionInterface(_session_dir)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize database and adapters
init_db()
payment_adapter = create_payment_adapter()  # Config-driven adapter selection
humanizer_adapter = RuleBasedHumanizer()

# Safety check: refuse to start in production with mock payment adapter
if app.config.get('PAYMENT_ADAPTER') == 'mock' and os.environ.get('FLASK_ENV') == 'production':
    raise RuntimeError(
        "Refusing to start: PAYMENT_ADAPTER=mock is not allowed in production. "
        "Set PAYMENT_ADAPTER=alipay and configure Alipay credentials."
    )

if app.config.get('PAYMENT_ADAPTER') == 'mock':
    import logging
    logging.warning("PAYMENT_ADAPTER=mock is enabled. This should only be used for development.")


# ========== Database Connection Helpers ==========

def get_db():
    """Get a database connection scoped to the current request context."""
    if 'db_conn' not in g:
        g.db_conn = get_connection()
    return g.db_conn


@app.teardown_appcontext
def close_db(exception=None):
    """Close the database connection at the end of each request."""
    conn = g.pop('db_conn', None)
    if conn is not None:
        conn.close()


# ========== Constants ==========
PRICE_PER_1000_WORDS = 14.9  # ¥14.9 / 1000 words
FREE_WORD_LIMIT = 500  # Words requiring payment for rewrite
MAX_FREE_ANALYSIS_WORDS = 500  # Max words for free analysis

# ========== Text Extraction ==========

def extract_text_from_docx(filepath):
    """Extract text from .docx file."""
    from docx import Document
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return '\n\n'.join(paragraphs)


def extract_text_from_pdf(filepath):
    """Extract text from .pdf file."""
    import fitz
    doc = fitz.open(filepath)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return '\n\n'.join(text_parts)


def extract_text(filepath):
    """Extract text from uploaded file based on extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_text_from_docx(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    elif ext == '.md':
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ========== Format Output Helpers ==========

def generate_docx(text):
    """Generate a .docx file in-memory from text content."""
    from docx import Document
    doc = Document()
    for paragraph in text.split('\n\n'):
        p = doc.add_paragraph(paragraph.strip())
        if not paragraph.strip():
            p.add_run(' ')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_pdf(text):
    """Generate a .pdf file in-memory from text content."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    # Insert text into the page
    rect = page.rect
    page.insert_text(
        (rect.x0 + 50, rect.y0 + 50),
        text,
        fontsize=11,
        fontname="helv",
        color=(0, 0, 0)
    )
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return buf


def generate_file_response(text, original_format, filename):
    """Generate a file response for download based on format."""
    base_name = os.path.splitext(filename)[0] if filename else 'humanized'

    if original_format == 'docx':
        buf = generate_docx(text)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{base_name}_humanized.docx'
        )
    elif original_format == 'pdf':
        buf = generate_pdf(text)
        return send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{base_name}_humanized.pdf'
        )
    elif original_format == 'md':
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/markdown',
            as_attachment=True,
            download_name=f'{base_name}_humanized.md'
        )
    else:  # txt (default)
        buf = io.BytesIO(text.encode('utf-8'))
        return send_file(
            buf,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{base_name}_humanized.txt'
        )


# ========== Decorator ==========

def login_required(f):
    """Require user to be logged in. Returns 401 with login_required flag."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('user_id'):
            return jsonify({"error": "请先登录", "login_required": True}), 401
        return f(*args, **kwargs)
    return decorated_function


# ========== Helper ==========

def _generate_modification_suggestions(analysis_result, text):
    """Generate suggestions based on AI analysis results."""
    suggestions = []
    sub_scores = analysis_result.get("sub_scores", {})
    sub_details = analysis_result.get("sub_score_details", {})

    # 1. Perplexity suggestion
    if sub_scores.get("perplexity_score", 0) > 50:
        suggestions.append({
            "target": "perplexity",
            "icon": "📊",
            "title": "词汇多样性不足",
            "detail": "你的文本词汇模式过于可预测，AI检测模型容易识别。建议增加同义词替换和句式变化。",
            "severity": "high" if sub_scores["perplexity_score"] > 70 else "medium"
        })

    # 2. Pattern suggestion
    pattern_data = sub_details.get("pattern", {})
    if pattern_data.get("ai_phrase_count", 0) > 3:
        top_phrases = pattern_data.get("top_phrases", [])
        suggestions.append({
            "target": "pattern",
            "icon": "🔍",
            "title": f"检测到 {pattern_data['ai_phrase_count']} 个AI常用短语",
            "detail": f"常见AI短语如「{'」、「'.join(top_phrases[:3])}」在AI生成文本中频繁出现，替换为更自然的表达可降低AI率。",
            "severity": "high"
        })

    # 3. Readability suggestion
    readability = sub_details.get("readability", {})
    fk_grade = readability.get("flesch_kincaid", 10)
    avg_sent = readability.get("avg_sentence_length", 20)
    if fk_grade > 14 or avg_sent > 25:
        suggestions.append({
            "target": "readability",
            "icon": "✂️",
            "title": f"句长过于均匀（平均 {avg_sent:.0f} 词/句）",
            "detail": "AI生成的文本句子长度变化较小，缺乏人类写作的自然节奏感。建议混合长短句，增加句长变化。",
            "severity": "high" if avg_sent > 30 else "medium"
        })

    # 4. Burstiness suggestion
    if sub_scores.get("burstiness_score", 50) < 30:
        suggestions.append({
            "target": "burstiness",
            "icon": "📏",
            "title": "句式变化不足",
            "detail": "句子长度和结构变化不够丰富。建议混入短句（<10词）和长句（>30词），打破AI写作的规律性。",
            "severity": "medium"
        })

    # 5. Structure suggestion
    structure = sub_details.get("structure", {})
    if structure.get("formulaic_ratio", 0) > 0.2:
        suggestions.append({
            "target": "structure",
            "icon": "🏗️",
            "title": "句式开头较为刻板",
            "detail": "过多句子以「It is」「This is」「There is」等固定模式开头，建议变化句子起始方式。",
            "severity": "medium"
        })

    # Default suggestion if nothing specific
    if not suggestions:
        suggestions.append({
            "target": "general",
            "icon": "✅",
            "title": "文本质量良好",
            "detail": "AI检测指标正常，当前文本不太可能被标记为AI生成。",
            "severity": "low"
        })

    return suggestions


# ========== Auth Routes ==========

@app.route('/api/register', methods=['POST'])
def api_register():
    """Register a new user account."""
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    confirm_password = data.get('confirm_password', '')

    # Validate email format
    if not email or '@' not in email or '.' not in email:
        return jsonify({"error": "请输入有效的邮箱地址"}), 400

    # Validate password length
    if len(password) <= 6:
        return jsonify({"error": "密码长度必须大于 6 位"}), 400

    # Validate password confirmation
    if password != confirm_password:
        return jsonify({"error": "两次密码输入不一致"}), 400

    conn = get_db()

    # Check if email already exists
    existing = User.get_by_email(conn, email)
    if existing:
        return jsonify({"error": "该邮箱已被注册"}), 409

    try:
        user = User.create(conn, email, password)
        session['user_id'] = user['id']
        return jsonify({
            "success": True,
            "user": {"id": user['id'], "email": user['email']}
        })
    except Exception as e:
        return jsonify({"error": f"注册失败：{str(e)}"}), 500


@app.route('/api/login', methods=['POST'])
def api_login():
    """Log in an existing user."""
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({"error": "请填写邮箱和密码"}), 400

    conn = get_db()
    user = User.verify_password(conn, email, password)
    if not user:
        return jsonify({"error": "邮箱或密码错误"}), 401

    session['user_id'] = user['id']
    return jsonify({
        "success": True,
        "user": {"id": user['id'], "email": user['email']}
    })


@app.route('/api/logout', methods=['POST'])
def api_logout():
    """Log out the current user."""
    session.pop('user_id', None)
    return jsonify({"success": True})


@app.route('/api/me')
def api_me():
    """Get current logged-in user info."""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "未登录"}), 401

    conn = get_db()
    user = User.get_by_id(conn, user_id)
    if not user:
        session.pop('user_id', None)
        return jsonify({"error": "未登录"}), 401

    return jsonify({
        "user": {"id": user['id'], "email": user['email']}
    })


# ========== Routes ==========

@app.route('/')
def index():
    """Landing page."""
    return render_template('index.html')


@app.route('/orders')
def orders_page():
    """Order history page — requires login."""
    user_id = session.get('user_id')
    if not user_id:
        return render_template('orders.html', needs_login=True)
    return render_template('orders.html', needs_login=False)


@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    """
    Analyze text for AI content.
    Accepts: text (direct paste) OR file (upload)
    Returns: AI score, paragraph analysis, suggestions
    """
    text = None
    filename = None
    original_format = 'txt'
    original_filename = None

    # Check if file was uploaded
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ['.docx', '.pdf', '.txt', '.md']:
                return jsonify({"error": "仅支持 .docx、.pdf、.txt、.md 格式"}), 400

            original_filename = file.filename
            original_format = ext[1:]  # Remove the dot
            filename = f"{uuid.uuid4().hex}{ext}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            text = extract_text(filepath)
            # Clean up the uploaded file after extraction
            try:
                os.remove(filepath)
            except OSError:
                pass

    # Check if text was pasted
    if not text:
        data = request.get_json(silent=True) or {}
        text = data.get('text', '').strip()
        if not text:
            return jsonify({"error": "请上传文档或粘贴英文文本"}), 400

    # Clean the text
    text = text.strip()

    # Store original format info in session
    session['last_original_format'] = original_format
    session['last_original_filename'] = original_filename
    session['last_text'] = text

    if len(text) < 50:
        return jsonify({"error": "文本太短，请提供至少 50 个字符"}), 400

    word_count = len(text.split())

    # Enforce free word limit
    if word_count > MAX_FREE_ANALYSIS_WORDS:
        return jsonify({
            "error": f"免费检测限制 {MAX_FREE_ANALYSIS_WORDS} 词以内（当前 {word_count} 词）",
            "over_limit": True,
            "word_count": word_count,
            "max_free_words": MAX_FREE_ANALYSIS_WORDS,
            "extracted_text": text,
            "price": round(max(PRICE_PER_1000_WORDS * (word_count / 1000), PRICE_PER_1000_WORDS), 2),
            "original_format": original_format,
            "original_filename": original_filename
        }), 413

    is_paid = word_count > FREE_WORD_LIMIT

    # Run AI detection
    try:
        full_analysis = analyze_text(text)
        paragraph_analysis = analyze_by_paragraphs(text)
    except Exception as e:
        return jsonify({"error": f"分析出错：{str(e)}"}), 500

    # Generate suggestions
    suggestions = _generate_modification_suggestions(full_analysis, text)

    # Calculate price
    price = max(PRICE_PER_1000_WORDS * (word_count / 1000), PRICE_PER_1000_WORDS)

    # Store in session for later use
    session['last_text'] = text
    session['last_word_count'] = word_count
    session['last_price'] = round(price, 2)

    return jsonify({
        "success": True,
        "analysis": {
            "overall": full_analysis,
            "paragraphs": paragraph_analysis,
            "suggestions": suggestions
        },
        "text_preview": text[:500] + "..." if len(text) > 500 else text,
        "word_count": word_count,
        "price": round(price, 2),
        "is_paid": is_paid,
        "extracted_text": text,
        "original_format": original_format,
        "original_filename": original_filename
    })


@app.route('/api/suggestion-detail', methods=['POST'])
def api_suggestion_detail():
    """
    Get detailed suggestions for a specific paragraph or section.
    """
    data = request.get_json(silent=True) or {}
    paragraph_text = data.get('text', '').strip()
    paragraph_index = data.get('paragraph_index', 0)

    if not paragraph_text or len(paragraph_text) < 50:
        return jsonify({"error": "段落文本太短"}), 400

    try:
        analysis = analyze_text(paragraph_text)
        suggestions = _generate_modification_suggestions(analysis, paragraph_text)
    except Exception as e:
        return jsonify({"error": f"分析出错：{str(e)}"}), 500

    return jsonify({
        "success": True,
        "analysis": analysis,
        "suggestions": suggestions,
        "paragraph_index": paragraph_index
    })


@app.route('/api/rewrite', methods=['POST'])
@login_required
def api_rewrite():
    """
    Rewrite text to reduce AI detection score.
    Simulates payment confirmation. Requires login.
    """
    data = request.get_json(silent=True) or {}
    text = data.get('text', session.get('last_text', ''))
    mode = data.get('mode', 'academic')  # 'academic' or 'aggressive'

    if not text:
        return jsonify({"error": "没有可改写的文本，请先分析"}), 400

    word_count = len(text.split())
    price = max(PRICE_PER_1000_WORDS * (word_count / 1000), PRICE_PER_1000_WORDS)

    # Generate a simulated order ID
    order_id = f"ORD-{uuid.uuid4().hex[:8].upper()}"

    # Store rewrite request for payment confirmation
    session['pending_rewrite'] = {
        'text': text,
        'mode': mode,
        'word_count': word_count,
        'price': round(price, 2),
        'order_id': order_id
    }

    return jsonify({
        "success": True,
        "order": {
            "order_id": order_id,
            "word_count": word_count,
            "price": round(price, 2),
            "mode": mode
        }
    })


@app.route('/api/confirm-payment', methods=['POST'])
@login_required
def api_confirm_payment():
    """
    Confirm payment and execute the rewrite.
    Requires payment_token from the frontend (simulated). Requires login.
    """
    pending = session.get('pending_rewrite')
    if not pending:
        return jsonify({"error": "没有待处理的改写请求"}), 400

    # Payment validation via adapter
    data = request.get_json(silent=True) or {}
    payment_token = data.get('payment_token', '')

    if not payment_adapter.verify_payment(payment_token):
        return jsonify({"error": "支付验证失败，请重新尝试"}), 402

    text = pending['text']
    mode = pending['mode']
    order_id = pending['order_id']

    try:
        # Run humanization via adapter
        humanized = humanizer_adapter.humanize(text, mode=mode)

        # Analyze original text
        original_analysis = analyze_text(text)
        rewritten_analysis = analyze_text(humanized)

        # Generate diff-like comparison
        original_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        rewritten_paragraphs = [p.strip() for p in humanized.split('\n\n') if p.strip()]

        paragraph_comparison = []
        for i, (orig_p, new_p) in enumerate(zip(original_paragraphs, rewritten_paragraphs)):
            if len(orig_p) >= 100 and len(new_p) >= 100:
                paragraph_comparison.append({
                    "index": i,
                    "original_preview": orig_p[:150] + "..." if len(orig_p) > 150 else orig_p,
                    "rewritten_preview": new_p[:150] + "..." if len(new_p) > 150 else new_p,
                    "original_score": round(original_analysis['ai_score'], 1),
                    "rewritten_score": round(rewritten_analysis['ai_score'], 1),
                    "reduction": round(original_analysis['ai_score'] - rewritten_analysis['ai_score'], 1)
                })

        # Save order if user is logged in
        user_id = session.get('user_id')
        if user_id:
            original_format = session.get('last_original_format', 'txt')
            original_filename = session.get('last_original_filename', None)
            try:
                conn = get_db()
                Order.create(
                    conn,
                    user_id=user_id,
                    order_id=order_id,
                    original_text=text,
                    rewritten_text=humanized,
                    original_format=original_format,
                    original_filename=original_filename,
                    word_count=pending['word_count'],
                    price=pending['price'],
                    mode=mode,
                    original_score=original_analysis.get('ai_score', 0),
                    rewritten_score=rewritten_analysis.get('ai_score', 0)
                )
            except Exception:
                # If saving fails, still return result (non-fatal)
                pass

        # Clean up session
        session.pop('pending_rewrite', None)
        session['last_rewritten'] = {
            'original': text,
            'rewritten': humanized,
            'original_score': original_analysis.get('ai_score', 0),
            'rewritten_score': rewritten_analysis.get('ai_score', 0),
            'order_id': order_id
        }

        return jsonify({
            "success": True,
            "order_id": order_id,
            "original": {
                "text": text,
                "ai_score": round(original_analysis['ai_score'], 1),
                "risk_level": original_analysis['risk_level']
            },
            "rewritten": {
                "text": humanized,
                "ai_score": round(rewritten_analysis['ai_score'], 1),
                "risk_level": rewritten_analysis['risk_level']
            },
            "improvement": round(original_analysis['ai_score'] - rewritten_analysis['ai_score'], 1),
            "paragraph_comparison": paragraph_comparison,
            "original_format": session.get('last_original_format', 'txt'),
            "original_filename": session.get('last_original_filename', None)
        })

    except Exception as e:
        return jsonify({"error": f"改写出错：{str(e)}"}), 500


@app.route('/api/preview-rewrite', methods=['POST'])
def api_preview_rewrite():
    """
    Preview what the rewritten text would look like (free preview, limited).
    """
    data = request.get_json(silent=True) or {}
    text = data.get('text', '')

    if not text:
        return jsonify({"error": "没有可预览改写的文本"}), 400

    word_count = len(text.split())

    # Only preview first paragraph if too long
    if word_count > 200:
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        if paragraphs:
            text = paragraphs[0]
        else:
            text = ' '.join(text.split()[:200])

    try:
        humanized = humanizer_adapter.humanize(text, mode='academic')
        original_analysis = analyze_text(text)
        rewritten_analysis = analyze_text(humanized)

        return jsonify({
            "success": True,
            "original_excerpt": text,
            "rewritten_excerpt": humanized,
            "original_score": round(original_analysis['ai_score'], 1),
            "rewritten_score": round(rewritten_analysis['ai_score'], 1),
            "note": "此为免费预览，仅展示部分内容。支付后可改写全文。"
        })

    except Exception as e:
        return jsonify({"error": f"预览出错：{str(e)}"}), 500


# ========== Payment API Routes ==========

@app.route('/api/create-payment', methods=['POST'])
@login_required
def api_create_payment():
    """
    Create a payment order and return QR code for scanning.
    This replaces the old /api/rewrite + /api/confirm-payment flow.
    """
    data = request.get_json(silent=True) or {}
    text = data.get('text', session.get('last_text', ''))
    mode = data.get('mode', 'academic')

    if not text:
        return jsonify({"error": "没有可改写的文本，请先分析"}), 400

    word_count = len(text.split())
    price = max(PRICE_PER_1000_WORDS * (word_count / 1000), PRICE_PER_1000_WORDS)
    price = round(price, 2)

    # Generate order ID
    order_id = f"ORD-{uuid.uuid4().hex[:8].upper()}"

    # Get format info from session
    original_format = session.get('last_original_format', 'txt')
    original_filename = session.get('last_original_filename', None)
    user_id = session.get('user_id')

    # Create pending payment record in DB
    conn = get_db()
    try:
        Order.create_payment_record(
            conn, user_id, order_id, text, original_format, original_filename,
            word_count, price, mode
        )
    except Exception as e:
        return jsonify({"error": f"创建订单失败：{str(e)}"}), 500

    # Call payment adapter to create prepay order
    result = payment_adapter.create_prepay_order(
        order_id, price, f"AI降AI率服务 - {word_count}词"
    )

    if result.get('error'):
        return jsonify({"error": result['error']}), 500

    # Save QR code to DB for reference
    qr_code = result.get('qr_code')
    if qr_code:
        Order.save_qr_code(conn, order_id, qr_code)

    return jsonify({
        "success": True,
        "order": {
            "order_id": order_id,
            "word_count": word_count,
            "price": price,
            "qr_code": qr_code,
            "mode": mode,
            "expires_in": result.get('expires_in', 1800)
        }
    })


@app.route('/api/payment-status/<order_id>')
def api_payment_status(order_id):
    """
    Check payment status for an order.
    Used by frontend polling after QR code is displayed.
    """
    user_id = session.get('user_id')
    conn = get_db()

    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return jsonify({"error": "订单不存在"}), 404

    # Check ownership (if logged in)
    if user_id and order['user_id'] != user_id:
        return jsonify({"error": "无权访问该订单"}), 403

    payment_status = order.get('payment_status', 'pending')
    status = order.get('status', 'pending')

    # If still pending, optionally query payment gateway directly
    if payment_status == 'pending':
        try:
            query_result = payment_adapter.query_payment(order_id)
            if query_result.get('status') == 'paid':
                # Payment detected! Mark as paid and trigger rewrite.
                _process_payment_success(conn, order_id, query_result.get('trade_no'))
                payment_status = 'paid'
                status = 'processing'
        except Exception:
            pass  # Query failed, just return current status

    # Build response based on status
    response = {
        "order_id": order_id,
        "payment_status": payment_status,
        "status": status,
        "price": order.get('price'),
        "word_count": order.get('word_count')
    }

    # If completed, include the rewrite result
    if status == 'completed' and order.get('rewritten_text'):
        original_analysis = analyze_text(order['original_text']) if order.get('original_text') else {}
        response.update({
            "success": True,
            "original": {
                "text": order['original_text'],
                "ai_score": round(order.get('original_score', 0), 1),
                "risk_level": original_analysis.get('risk_level', 'unknown')
            },
            "rewritten": {
                "text": order['rewritten_text'],
                "ai_score": round(order.get('rewritten_score', 0), 1),
                "risk_level": analyze_text(order['rewritten_text']).get('risk_level', 'unknown') if order['rewritten_text'] else 'unknown'
            },
            "improvement": round((order.get('original_score', 0) or 0) - (order.get('rewritten_score', 0) or 0), 1),
            "original_format": order.get('original_format', 'txt'),
            "original_filename": order.get('original_filename')
        })

    return jsonify(response)


@app.route('/api/webhook/alipay', methods=['POST'])
def api_webhook_alipay():
    """
    Alipay async notification webhook.
    This is called by Alipay servers after payment is completed.
    """
    # Get parameters (Alipay sends as form data)
    params = request.form.to_dict()

    sign = params.pop('sign', None)
    sign_type = params.pop('sign_type', None)

    # Verify notification
    is_valid, order_id, trade_no, amount = payment_adapter.verify_notification(params, sign)

    if not is_valid:
        return "fail", 200

    conn = get_db()

    # Check order exists and is pending
    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return "fail", 200

    if order.get('payment_status') != 'pending':
        # Already processed (duplicate notification)
        return "success", 200

    # Verify amount matches
    if amount and abs(amount - (order.get('price') or 0)) > 0.01:
        # Amount mismatch - potential fraud
        return "fail", 200

    # Process payment success
    try:
        _process_payment_success(conn, order_id, trade_no)
    except Exception as e:
        return "fail", 200

    return "success", 200


def _process_payment_success(conn, order_id, trade_no):
    """
    Internal function to handle successful payment.
    Marks order as paid and triggers rewrite in background.
    """
    import threading

    order = Order.get_by_order_id(conn, order_id)
    if not order:
        raise ValueError(f"Order {order_id} not found")

    # Mark as paid
    Order.mark_paid(conn, order_id, trade_no, datetime.utcnow().isoformat())

    # Read mode from DB — cannot access Flask session in background thread
    mode = order.get('mode', 'academic')

    # Trigger rewrite in background thread (don't block the webhook response)
    def do_rewrite():
        try:
            text = order['original_text']
            humanized = humanizer_adapter.humanize(text, mode=mode)
            rewritten_analysis = analyze_text(humanized)
            original_analysis = analyze_text(text)

            # Update order with result
            conn2 = get_connection()
            try:
                Order.update_result(
                    conn2, order_id, humanized,
                    rewritten_analysis.get('ai_score', 0),
                    original_analysis.get('ai_score', 0)
                )
            finally:
                conn2.close()
        except Exception as e:
            import logging
            logging.exception(f"Background rewrite failed for {order_id}")

    # Start background thread
    thread = threading.Thread(target=do_rewrite)
    thread.daemon = True
    thread.start()


# ========== Test/Debug API (only available in mock mode) ==========

@app.route('/api/test/mock-payment/<order_id>', methods=['POST'])
def api_test_mock_payment(order_id):
    """
    Simulate a successful payment for testing purposes.
    Only available when PAYMENT_ADAPTER=mock.
    """
    if app.config.get('PAYMENT_ADAPTER') != 'mock':
        return jsonify({"error": "仅在 mock 模式下可用"}), 403

    conn = get_db()
    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return jsonify({"error": "订单不存在"}), 404

    if order.get('payment_status') != 'pending':
        return jsonify({"error": f"订单状态不是 pending，当前: {order.get('payment_status')}"}), 400

    # Simulate payment success
    trade_no = f"MOCK_TRADE_{order_id}"
    try:
        _process_payment_success(conn, order_id, trade_no)
        return jsonify({"success": True, "message": "支付模拟成功，正在后台改写...", "order_id": order_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ========== Download API ==========

@app.route('/api/download/<order_id>')
def api_download(order_id):
    """
    Download rewritten text in the specified format.
    Requires login (or the order must belong to the current user's session).
    Query params: ?format=docx|pdf|txt|md (default: original_format)
    """
    user_id = session.get('user_id')
    conn = get_db()

    # Try to find the order
    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return jsonify({"error": "订单不存在"}), 404

    # Check ownership
    if user_id and order['user_id'] != user_id:
        return jsonify({"error": "无权访问该订单"}), 403

    # If not logged in, check if this is the last rewritten order in session
    if not user_id:
        last = session.get('last_rewritten', {})
        if last.get('order_id') != order_id:
            return jsonify({"error": "请登录后下载"}), 401

    # Determine output format
    req_format = request.args.get('format', order.get('original_format', 'txt'))
    if req_format not in ['docx', 'pdf', 'txt', 'md']:
        req_format = order.get('original_format', 'txt')

    # Generate file response
    rewritten_text = order['rewritten_text']
    filename = order.get('original_filename', 'humanized')
    return generate_file_response(rewritten_text, req_format, filename)


# ========== Order API Routes ==========

@app.route('/api/orders')
def api_orders():
    """Get user's order list with pagination. Requires login."""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "未登录"}), 401

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)

    conn = get_db()
    orders, total = Order.get_by_user_id(conn, user_id, page=page, per_page=per_page)

    total_pages = max(1, (total + per_page - 1) // per_page)

    return jsonify({
        "orders": orders,
        "total": total,
        "page": page,
        "pages": total_pages
    })


@app.route('/api/orders/<order_id>')
def api_order_detail(order_id):
    """Get details for a specific order. Requires login."""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "未登录"}), 401

    conn = get_db()
    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return jsonify({"error": "订单不存在"}), 404

    if order['user_id'] != user_id:
        return jsonify({"error": "无权访问该订单"}), 403

    return jsonify({"order": order})


@app.route('/api/orders/<order_id>/rehumanize', methods=['POST'])
def api_rehumanize(order_id):
    """
    Re-humanize an existing order (free within 7 days).
    Requires login and non-expired order.
    """
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "未登录"}), 401

    data = request.get_json(silent=True) or {}
    mode = data.get('mode', 'academic')

    conn = get_db()
    order = Order.get_by_order_id(conn, order_id)
    if not order:
        return jsonify({"error": "订单不存在"}), 404

    if order['user_id'] != user_id:
        return jsonify({"error": "无权操作该订单"}), 403

    # Check expiration
    expires_at = order['expires_at']
    try:
        expires_dt = datetime.fromisoformat(expires_at)
    except (ValueError, TypeError):
        return jsonify({"error": "订单日期异常"}), 500

    if datetime.utcnow() > expires_dt:
        return jsonify({"error": "订单已过期（超过 7 天），请重新购买"}), 410

    try:
        # Re-humanize
        original_text = order['original_text']
        humanized = humanizer_adapter.humanize(original_text, mode=mode)
        rewritten_analysis = analyze_text(humanized)

        # Update the order
        Order.update_rewrite(conn, order_id, humanized, rewritten_analysis.get('ai_score', 0))

        original_score = order.get('original_score', 0)

        return jsonify({
            "success": True,
            "order_id": order_id,
            "original": {
                "text": original_text,
                "ai_score": round(original_score, 1)
            },
            "rewritten": {
                "text": humanized,
                "ai_score": round(rewritten_analysis['ai_score'], 1),
                "risk_level": rewritten_analysis['risk_level']
            },
            "improvement": round(original_score - rewritten_analysis['ai_score'], 1)
        })

    except Exception as e:
        return jsonify({"error": f"改写出错：{str(e)}"}), 500


# ========== Main ==========

if __name__ == '__main__':
    print("=" * 50)
    print("AI Humanizer - Starting on http://127.0.0.1:5100")
    print("=" * 50)
    app.run(host='127.0.0.1', port=5100)
